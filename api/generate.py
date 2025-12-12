from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from http.server import BaseHTTPRequestHandler
import io
import json
from datetime import datetime
import re
import os
import urllib.request
import urllib.parse

# Russian number spelling dictionaries
ONES = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
ONES_FEM = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
TENS = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
TEENS = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать', 'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
HUNDREDS = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
THOUSANDS = ['', 'тысяча', 'тысячи', 'тысяч']
MILLIONS = ['', 'миллион', 'миллиона', 'миллионов']
RUBLES = ['рубль', 'рубля', 'рублей']
KOPECKS = ['копейка', 'копейки', 'копеек']
WORKDAYS = ['рабочий день', 'рабочих дня', 'рабочих дней']

MONTHS_RU = ['', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']

def format_number_russian(num):
    """Format number with Russian formatting: space as thousands separator, no decimals"""
    # Round to integer (no cents/kopecks)
    num_int = int(round(num))
    
    # Add space as thousands separator
    integer_formatted = ''
    for i, digit in enumerate(reversed(str(num_int))):
        if i > 0 and i % 3 == 0:
            integer_formatted = ' ' + integer_formatted
        integer_formatted = digit + integer_formatted
    
    return integer_formatted

def get_declension(num, forms):
    """Get proper Russian declension based on number"""
    num = abs(num) % 100
    num1 = num % 10
    if num > 10 and num < 20:
        return forms[2]
    if num1 > 1 and num1 < 5:
        return forms[1]
    if num1 == 1:
        return forms[0]
    return forms[2]

def spell_number_russian(num, feminine=False):
    """Convert number to Russian words"""
    if num == 0:
        return 'ноль'
    
    result = []
    num = int(num)
    
    # Millions
    if num >= 1000000:
        millions = num // 1000000
        result.append(spell_number_russian(millions, False))
        result.append(get_declension(millions, MILLIONS))
        num = num % 1000000
    
    # Thousands
    if num >= 1000:
        thousands = num // 1000
        if thousands == 1:
            result.append('одна')
        elif thousands == 2:
            result.append('две')
        else:
            result.append(spell_number_russian(thousands, True))
        result.append(get_declension(thousands, THOUSANDS))
        num = num % 1000
    
    # Hundreds
    if num >= 100:
        hundreds = num // 100
        result.append(HUNDREDS[hundreds])
        num = num % 100
    
    # Tens and ones
    if num >= 20:
        tens = num // 10
        result.append(TENS[tens])
        num = num % 10
    
    if num >= 10:
        result.append(TEENS[num - 10])
        num = 0
    
    if num > 0:
        if feminine:
            result.append(ONES_FEM[num])
        else:
            result.append(ONES[num])
    
    return ' '.join(result)

def spell_money_russian(amount):
    """Convert money amount to Russian words (no kopecks)"""
    # Round to integer (no cents/kopecks)
    rubles = int(round(amount))
    
    rubles_text = spell_number_russian(rubles, False)
    rubles_form = get_declension(rubles, RUBLES)
    
    result = f"{rubles_text.capitalize()} {rubles_form}"
    
    return result

def spell_workdays_russian(days):
    """Convert workdays number to Russian words"""
    days_text = spell_number_russian(days, False)
    days_form = get_declension(days, WORKDAYS)
    return f"{days_text.capitalize()} {days_form}"

def format_date_full_russian(date_str):
    """Format date to full Russian format"""
    try:
        # Try parsing different date formats
        if ' ' in date_str:
            # Format like "Jan 15, 2025"
            dt = datetime.strptime(date_str, '%b %d, %Y')
        elif '-' in date_str:
            # Format like "2025-01-15"
            dt = datetime.strptime(date_str.split()[0], '%Y-%m-%d')
        else:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
        
        day = dt.day
        month = MONTHS_RU[dt.month]
        year = dt.year
        
        return f"{day} {month} {year} г."
    except:
        return date_str

def set_cell_border(cell, **kwargs):
    """Set border for a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data['val'])
            element.set(qn('w:sz'), str(edge_data.get('sz', 6)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            borders.append(element)
    tcPr.append(borders)

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def set_font_times_new_roman(run, size=12, bold=False, italic=False):
    """Set font to Times New Roman with specified size and style"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_company_header(doc, company_data, locale='ru-RU', available_width=Mm(190)):
    """Add company information header in Russian format using a table with border"""
    # Create header table
    header_table = doc.add_table(rows=1, cols=1)
    header_table.columns[0].width = available_width
    
    header_cell = header_table.cell(0, 0)
    
    # Line 1: ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ
    line1_para = header_cell.paragraphs[0]
    line1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1_run = line1_para.add_run("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ")
    set_font_times_new_roman(line1_run, size=12, bold=True, italic=True)
    
    # Line 2: «МК СЕРВИС»
    line2_para = header_cell.add_paragraph()
    line2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_name = company_data.get('name', '')
    line2_run = line2_para.add_run(f"«{company_name}»")
    set_font_times_new_roman(line2_run, size=24, bold=True, italic=True)
    
    # Line 3: Address
    if company_data.get('address'):
        line3_para = header_cell.add_paragraph()
        line3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line3_run = line3_para.add_run(company_data['address'])
        set_font_times_new_roman(line3_run, size=9, bold=False, italic=False)
    
    # Line 4: Bank details
    if locale and locale.startswith('ru'):
        bank_parts = []
        if company_data.get('bankName'):
            bank_parts.append(company_data['bankName'])
        if company_data.get('bankAccount'):
            bank_parts.append(f"р/с {company_data['bankAccount']}")
        if company_data.get('correspondentAccount'):
            bank_parts.append(f"к/с {company_data['correspondentAccount']}")
        if company_data.get('bankBik'):
            bank_parts.append(f"БИК {company_data['bankBik']}")
        
        if bank_parts:
            line4_para = header_cell.add_paragraph()
            line4_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line4_run = line4_para.add_run(', '.join(bank_parts))
            set_font_times_new_roman(line4_run, size=9, bold=False, italic=False)
    
    # Line 5: ИНН and КПП
    if locale and locale.startswith('ru'):
        company_parts = []
        if company_data.get('inn'):
            company_parts.append(f"ИНН {company_data['inn']}")
        if company_data.get('kpp'):
            company_parts.append(f"КПП {company_data['kpp']}")
        
        if company_parts:
            line5_para = header_cell.add_paragraph()
            line5_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line5_run = line5_para.add_run(', '.join(company_parts))
            set_font_times_new_roman(line5_run, size=9, bold=False, italic=False)
    
    # Set bottom border for the table cell
    set_cell_border(header_cell, bottom={'val': 'single', 'sz': 6, 'color': '000000'})
    
    doc.add_paragraph()

def extract_order_number(order_id):
    """Extract number from order ID (e.g., 'order-27193' -> '27193', 'KP-2025-193' -> '193')"""
    if not order_id:
        return ''
    order_str = str(order_id)
    
    # If order ID contains hyphens, get the last part after the last hyphen
    if '-' in order_str:
        parts = order_str.split('-')
        last_part = parts[-1]
        # Extract digits from the last part only
        digits_only = re.sub(r'\D', '', last_part)
        if digits_only:
            return digits_only
    
    # If no hyphens, extract all digits
    digits_only = re.sub(r'\D', '', order_str)
    if digits_only:
        return digits_only
    
    return order_str

def add_document_header(doc, order_data, client_data, doc_type, locale='ru-RU', available_width=Mm(190)):
    """Add document header (Смета №, КП №, or Спецификация №)"""
    if doc_type == 'invoice':
        doc_label = "Смета №"
    elif doc_type == 'specification':
        doc_label = "Спецификация №"
    else:  # po
        doc_label = "КП №"
    
    # Always use order ID to extract number
    order_id = order_data.get('id', '')
    doc_number = extract_order_number(order_id)
    
    # Use current date
    current_date = datetime.now()
    if locale and locale.startswith('ru'):
        date_formatted = format_date_full_russian(current_date.strftime('%Y-%m-%d'))
    else:
        date_formatted = current_date.strftime('%Y-%m-%d')
    
    # Create table for alignment - use full width
    header_table = doc.add_table(rows=2, cols=2)
    # Extract mm value from Mm object
    width_val = available_width.mm if hasattr(available_width, 'mm') else float(available_width)
    header_table.columns[0].width = Mm(width_val * 0.65)  # 65% for left content
    header_table.columns[1].width = Mm(width_val * 0.35)  # 35% for date
    
    # Left column: Document number and client
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(f"{doc_label} {doc_number}")
    set_font_times_new_roman(left_run, size=12, bold=False, italic=False)
    
    # Client name on next line
    client_name = client_data.get('company') or client_data.get('name') or 'Клиента'
    left_para2 = left_cell.add_paragraph()
    left_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run2 = left_para2.add_run(f"для {client_name}")
    set_font_times_new_roman(left_run2, size=12, bold=False, italic=False)
    
    # Right column: Date (right-aligned, spanning both rows)
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(date_formatted)
    set_font_times_new_roman(right_run, size=12, bold=False, italic=False)
    
    # Merge right cell to span both rows
    right_cell2 = header_table.cell(1, 1)
    right_cell.merge(right_cell2)
    
    doc.add_paragraph()
    
    # Order title (centered, not bold)
    if order_data.get('orderTitle'):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(order_data['orderTitle'])
        set_font_times_new_roman(title_run, size=12, bold=False, italic=False)
    
    doc.add_paragraph()

def calculate_column_widths(jobs_data, available_width):
    """
    Calculate column widths with explicit fixed widths for consistent rendering
    across Word Online and Word Desktop.
    Returns a list of widths in mm for each column.
    """
    width_val = available_width.mm if hasattr(available_width, 'mm') else float(available_width)
    
    # Fixed widths for columns (in millimeters)
    # Using explicit widths in mm ensures consistent rendering across regional Word settings
    fixed_widths = [
        10,  # Column 0: № (number)
        0,   # Column 1: Наименование (name) - calculated as remaining space
        20,  # Column 2: Кол-во (quantity)
        27,  # Column 3: Стоимость (unit price)
        27,  # Column 4: Сумма (total)
    ]
    
    # Calculate remaining space for Наименование column
    fixed_total = sum(fixed_widths)
    remaining_space = width_val - fixed_total
    
    # Ensure minimum width for Наименование (at least 50 mm)
    if remaining_space < 50:
        remaining_space = 50
    
    fixed_widths[1] = remaining_space
    
    return fixed_widths

def add_work_description(doc, jobs_data, order_data, doc_type='invoice', locale='ru-RU', work_days=30, available_width=Mm(190)):
    """Add work description with table view for jobs"""
    if not jobs_data:
        doc.add_paragraph("Нет работ")
        return
    
    # Create jobs table with 5 columns - use full available width
    jobs_table = doc.add_table(rows=1, cols=5)
    jobs_table.style = 'Table Grid'
    
    # Set table preferred width in mm to ensure Word uses metric units
    width_val = available_width.mm if hasattr(available_width, 'mm') else float(available_width)
    jobs_table.autofit = False
    tbl = jobs_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Set table width explicitly in mm (dxa units: 1mm = 56.7 twips, 1 inch = 1440 twips = 25.4mm)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_val * 56.7)))  # Convert mm to twips (dxa)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Calculate column widths based on content
    column_widths = calculate_column_widths(jobs_data, available_width)
    for i, width in enumerate(column_widths):
        col = jobs_table.columns[i]
        col.width = Mm(width)
        # Explicitly set width type to dxa (twips) to ensure Word uses absolute measurements
        # 1mm = 56.7 twips (dxa units)
        for cell in col.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width * 56.7)))  # Convert mm to twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    
    # Set minimum row height for all rows (8mm ≈ 0.3 inches)
    jobs_table.rows[0].height = Mm(8)
    
    # Header row
    header_cells = jobs_table.rows[0].cells
    headers = ['№', 'Наименование', 'Кол-во', 'Стоимость', 'Сумма']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Set grey background
        set_cell_shading(cell, 'D3D3D3')
        # Set vertical alignment to center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font_times_new_roman(run, size=12, bold=True, italic=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add job rows
    for idx, job in enumerate(jobs_data, start=1):
        new_row = jobs_table.add_row()
        new_row.height = Mm(8)
        row_cells = new_row.cells
        
        # №
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run0 = row_cells[0].paragraphs[0].add_run(str(idx))
        set_font_times_new_roman(run0, size=12, bold=False, italic=False)
        
        # Наименование
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = row_cells[1].paragraphs[0].add_run(job.get('name', ''))
        set_font_times_new_roman(run1, size=12, bold=False, italic=False)
        
        # Кол-во (without unit)
        job_qty = job.get('qty', '0')
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = row_cells[2].paragraphs[0].add_run(job_qty)
        set_font_times_new_roman(run2, size=12, bold=False, italic=False)
        
        # Цена за единицу (after markup) - calculate from lineTotal / quantity
        job_total = float(job.get('lineTotal', '0.00'))
        job_qty_float = float(job.get('qty', '1'))
        job_price_after_markup = job_total / job_qty_float if job_qty_float > 0 else 0.00
        row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run3 = row_cells[3].paragraphs[0].add_run(format_number_russian(job_price_after_markup))
        set_font_times_new_roman(run3, size=12, bold=False, italic=False)
        
        # Стоимость
        row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run4 = row_cells[4].paragraphs[0].add_run(format_number_russian(job_total))
        set_font_times_new_roman(run4, size=12, bold=False, italic=False)
    
    # Add subtotal row
    subtotal_table_row = jobs_table.add_row()
    subtotal_table_row.height = Mm(8)
    subtotal_row = subtotal_table_row.cells
    
    # Empty cells for columns 0-2
    for i in range(3):
        subtotal_row[i].text = ''
        subtotal_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Column 3: "Итого:" label
    subtotal_row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    subtotal_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtotal_run = subtotal_row[3].paragraphs[0].add_run('Итого:')
    set_font_times_new_roman(subtotal_run, size=12, bold=True, italic=False)
    
    # Column 4: Total sum
    total_sum = sum(float(job.get('lineTotal', '0.00')) for job in jobs_data)
    subtotal_row[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    subtotal_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_run = subtotal_row[4].paragraphs[0].add_run(format_number_russian(total_sum))
    set_font_times_new_roman(total_run, size=12, bold=True, italic=False)
    
    doc.add_paragraph()
    
    # Cost line with Russian number formatting and spelling on same line
    total = float(order_data.get('total', '0.00'))
    cost_para = doc.add_paragraph()
    if locale and locale.startswith('ru'):
        total_spelled = spell_money_russian(total)
        # Text before amount (regular) - different for specification
        if doc_type == 'specification':
            cost_text = "Стоимость поставки составляет: "
        else:
            cost_text = "Стоимость работ по заказу составляет: "
        cost_run1 = cost_para.add_run(cost_text)
        set_font_times_new_roman(cost_run1, size=12, bold=False, italic=False)
        # Amount (not bold)
        cost_run2 = cost_para.add_run(f"{format_number_russian(total)} руб.")
        set_font_times_new_roman(cost_run2, size=12, bold=False, italic=False)
        # Text after amount (regular)
        cost_run3 = cost_para.add_run(f" ({total_spelled}) Без НДС.")
        set_font_times_new_roman(cost_run3, size=12, bold=False, italic=False)
    else:
        # Text before amount (regular) - different for specification
        if doc_type == 'specification':
            cost_text = "Стоимость поставки составляет: "
        else:
            cost_text = "Стоимость работ по заказу составляет: "
        cost_run1 = cost_para.add_run(cost_text)
        set_font_times_new_roman(cost_run1, size=12, bold=False, italic=False)
        # Amount (not bold)
        cost_run2 = cost_para.add_run(f"{format_number_russian(total)} руб.")
        set_font_times_new_roman(cost_run2, size=12, bold=False, italic=False)
    
    # Add new line after cost line
    doc.add_paragraph()
    
    # Work completion deadline / Delivery deadline
    if locale and locale.startswith('ru'):
        days_spelled = spell_workdays_russian(work_days)
        deadline_para = doc.add_paragraph()
        if doc_type == 'specification':
            deadline_text = f"Срок поставки – {work_days} ({days_spelled}) рабочих дней с момента подписания спецификации и внесения предоплаты."
        else:
            deadline_text = f"Срок выполнения работ – {work_days} ({days_spelled}) рабочих дней с момента внесения предоплаты и подписания сметы."
        deadline_run = deadline_para.add_run(deadline_text)
        set_font_times_new_roman(deadline_run, size=12, bold=False, italic=False)
        
        # Note about deadline extension / Delivery location
        note_para = doc.add_paragraph()
        if doc_type == 'specification':
            note_text = "Место поставки – склад поставщика."
        else:
            note_text = "Сроки могут быть увеличены по согласованию сторон, в случае проведения дополнительных работ."
        note_run = note_para.add_run(note_text)
        set_font_times_new_roman(note_run, size=12, bold=False, italic=False)
        
        # Add 2 new lines after note
        doc.add_paragraph()
        doc.add_paragraph()
    
    doc.add_paragraph()

def add_footer_section(doc, company_data, doc_type, locale='ru-RU', available_width=Mm(190)):
    """Add footer section with signature and contact info"""
    if not (locale and locale.startswith('ru')):
        return
    
    doc.add_paragraph()
    
    if doc_type == 'invoice' or doc_type == 'specification':
        # For invoices and specifications: 2-column layout with signatures only
        footer_table = doc.add_table(rows=1, cols=2)
        # Extract mm value from Mm object
        width_val = available_width.mm if hasattr(available_width, 'mm') else float(available_width)
        footer_table.columns[0].width = Mm(width_val * 0.5)  # Исполнитель
        footer_table.columns[1].width = Mm(width_val * 0.5)  # Заказчик
        
        # Left column: Исполнитель
        left_cell = footer_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_header_run = left_para.add_run("Исполнитель")
        set_font_times_new_roman(left_header_run, size=12, bold=True, italic=False)
        
        # Double new line
        left_para.add_run("\n\n")
        
        # Signature lines: ________________/_________________/
        left_sig_run = left_para.add_run("________________/_________________/")
        set_font_times_new_roman(left_sig_run, size=12, bold=False, italic=False)
        
        # Right column: Заказчик
        right_cell = footer_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_header_run = right_para.add_run("Заказчик")
        set_font_times_new_roman(right_header_run, size=12, bold=True, italic=False)
        
        # Double new line
        right_para.add_run("\n\n")
        
        # Signature lines: ________________/_________________/
        right_sig_run = right_para.add_run("________________/_________________/")
        set_font_times_new_roman(right_sig_run, size=12, bold=False, italic=False)
    else:
        # For PO: 2-column layout - left: director info, right: contact info - use full width
        footer_table = doc.add_table(rows=1, cols=2)
        # Extract mm value from Mm object
        width_val = available_width.mm if hasattr(available_width, 'mm') else float(available_width)
        footer_table.columns[0].width = Mm(width_val * 0.55)  # Director info (55%)
        footer_table.columns[1].width = Mm(width_val * 0.45)  # Contact info (45%)
        
        # Left column: Director info
        left_cell = footer_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run1 = left_para.add_run("С уважением,")
        set_font_times_new_roman(left_run1, size=12, bold=False, italic=False)
        if company_data.get('directorName'):
            left_run2 = left_para.add_run("\n\nТехнический директор")
            set_font_times_new_roman(left_run2, size=12, bold=False, italic=False)
            left_run3 = left_para.add_run(f"\n{company_data.get('legalName', company_data.get('name', ''))}")
            set_font_times_new_roman(left_run3, size=12, bold=False, italic=False)
            left_run4 = left_para.add_run(f"\n{company_data['directorName']}")
            set_font_times_new_roman(left_run4, size=12, bold=False, italic=False)
        
        # Right column: Contact info (bottom right)
        right_cell = footer_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if company_data.get('email'):
            right_run1 = right_para.add_run(f"Email: {company_data['email']}")
            set_font_times_new_roman(right_run1, size=12, bold=False, italic=False)
        if company_data.get('phone'):
            if company_data.get('email'):
                right_run2 = right_para.add_run(f"\nТел: {company_data['phone']}")
            else:
                right_run2 = right_para.add_run(f"Тел: {company_data['phone']}")
            set_font_times_new_roman(right_run2, size=12, bold=False, italic=False)

def generate_document(data, doc_type):
    """Generate DOCX document from data"""
    doc = Document()
    
    # Set narrow margins (12.7mm ≈ 0.5 inches on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.left_margin = Mm(12.7)
        section.right_margin = Mm(12.7)
    
    # Available width with narrow margins: 215.9mm (A4) - 12.7mm - 12.7mm = 190.5mm
    AVAILABLE_WIDTH = Mm(190)
    
    # Set default font to Times New Roman, size 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set line spacing to 0 (single line spacing)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0  # Single line spacing (0 would be no spacing)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    locale = data.get('locale', 'ru-RU')
    company_data = data.get('company', {})
    client_data = data.get('client', {})
    order_data = data.get('order', {})
    jobs_data = data.get('jobs', [])
    work_days = data.get('workCompletionDays', 30)
    
    # 1. Company header with banking
    add_company_header(doc, company_data, locale, AVAILABLE_WIDTH)
    
    # 2. Document header (Смета № or КП №)
    add_document_header(doc, order_data, client_data, doc_type, locale, AVAILABLE_WIDTH)
    
    # 3. Work description with table view
    add_work_description(doc, jobs_data, order_data, doc_type, locale, work_days, AVAILABLE_WIDTH)
    
    # 4. Footer with signature and contact (3-column layout)
    add_footer_section(doc, company_data, doc_type, locale, AVAILABLE_WIDTH)
    
    return doc

# CORS Configuration
# Production domain
PRODUCTION_DOMAIN = 'https://client-code-one.vercel.app'

# Allowed origins list
ALLOWED_ORIGINS = [
    PRODUCTION_DOMAIN,
    PRODUCTION_DOMAIN.rstrip('/'),  # With and without trailing slash
]

# For local development, allow localhost
# Check if we're in production (Vercel sets VERCEL_ENV)
if os.environ.get('VERCEL_ENV') != 'production':
    ALLOWED_ORIGINS.extend([
        'http://localhost:3000',
        'http://localhost:5173',
        'http://127.0.0.1:3000',
        'http://127.0.0.1:5173',
    ])

def normalize_origin(origin):
    """Normalize origin by removing trailing slash"""
    if not origin:
        return None
    return origin.rstrip('/')

def is_origin_allowed(origin):
    """Check if origin is in allowed list"""
    if not origin:
        return False
    normalized = normalize_origin(origin)
    normalized_allowed = [normalize_origin(o) for o in ALLOWED_ORIGINS]
    return normalized in normalized_allowed

def verify_supabase_token(token):
    """
    Verify Supabase authentication token by calling Supabase auth API
    Returns user data if token is valid, None otherwise
    """
    if not token:
        return None
    
    try:
        supabase_url = os.environ.get('VITE_SUPABASE_URL')
        supabase_anon_key = os.environ.get('VITE_SUPABASE_ANON_KEY')
        
        if not supabase_url or not supabase_anon_key:
            # In production, these should be set, but if not, we can't verify
            # For local dev, we might not have them, so allow if token exists
            if os.environ.get('VERCEL_ENV') != 'production':
                return {'id': 'local-dev-user'}  # Allow in local dev
            return None
        
        # Call Supabase auth API to verify token
        # GET /auth/v1/user with Authorization header
        url = f"{supabase_url.rstrip('/')}/auth/v1/user"
        req = urllib.request.Request(url)
        req.add_header('Authorization', f'Bearer {token}')
        req.add_header('apikey', supabase_anon_key)
        
        with urllib.request.urlopen(req, timeout=5) as response:
            if response.status == 200:
                user_data = json.loads(response.read().decode('utf-8'))
                return user_data
            else:
                return None
    except Exception as e:
        # If verification fails for any reason, return None
        # Log error in development for debugging
        if os.environ.get('VERCEL_ENV') != 'production':
            print(f"Token verification error: {str(e)}")
        return None

class handler(BaseHTTPRequestHandler):
    def get_cors_headers(self):
        """Get CORS headers based on request origin"""
        origin = self.headers.get('Origin')
        headers = {}
        
        if origin and is_origin_allowed(origin):
            headers['Access-Control-Allow-Origin'] = origin
            headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS, GET'
            headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
            headers['Access-Control-Allow-Credentials'] = 'true'
        
        return headers
    
    def send_cors_headers(self):
        """Send CORS headers to response"""
        cors_headers = self.get_cors_headers()
        for key, value in cors_headers.items():
            self.send_header(key, value)
    
    def do_OPTIONS(self):
        """Handle CORS preflight requests"""
        self.send_response(200)
        self.send_cors_headers()
        self.send_header('Access-Control-Max-Age', '3600')
        self.end_headers()
    
    def do_POST(self):
        """Handle POST requests to generate DOCX documents"""
        try:
            print(f"POST request received at {self.path}")
            
            # Verify authentication
            auth_header = self.headers.get('Authorization')
            if not auth_header or not auth_header.startswith('Bearer '):
                self.send_error_response(401, 'Unauthorized: Authentication required')
                return
            
            token = auth_header.replace('Bearer ', '').strip()
            print(f"Token received: {token[:20]}...")
            
            # Check environment variables
            supabase_url = os.environ.get('VITE_SUPABASE_URL')
            supabase_key = os.environ.get('VITE_SUPABASE_ANON_KEY')
            print(f"Supabase URL set: {bool(supabase_url)}")
            print(f"Supabase key set: {bool(supabase_key)}")
            
            user = verify_supabase_token(token)
            
            if not user:
                self.send_error_response(401, 'Unauthorized: Invalid or expired token', 
                                       f'Supabase URL: {bool(supabase_url)}, Key: {bool(supabase_key)}')
                return
            
            print(f"User authenticated: {user.get('id', 'unknown')}")
            
            # Read request body
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length == 0:
                self.send_error_response(400, 'Request body is required')
                return
            
            body = self.rfile.read(content_length)
            data = json.loads(body.decode('utf-8'))
            doc_type = data.get('type', 'po')
            print(f"Generating document type: {doc_type}")
            
            # Generate document
            doc = generate_document(data, doc_type)
            print("Document generated successfully")
            
            # Save to BytesIO
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            doc_bytes = doc_buffer.getvalue()
            print(f"Document size: {len(doc_bytes)} bytes")
            
            # Generate filename using prefix from data and extract numbers from order ID
            document_prefix = data.get('documentPrefix', doc_type)
            order_id = data.get('order', {}).get('id', 'document')
            # Extract numbers from order ID (e.g., 'order-22650' -> '22650')
            order_numbers = re.sub(r'\D', '', str(order_id))
            filename = f"{document_prefix}-{order_numbers}.docx"
            
            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.send_cors_headers()
            self.send_header('Content-Length', str(len(doc_bytes)))
            self.end_headers()
            self.wfile.write(doc_bytes)
            print("Response sent successfully")
            
        except json.JSONDecodeError as e:
            self.send_error_response(400, f'Invalid JSON: {str(e)}', e)
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"Exception occurred: {error_trace}")
            self.send_error_response(500, f'Error generating document: {str(e)}', error_trace)
    
    def do_GET(self):
        """Handle GET requests for health check"""
        if self.path == '/api/generate' or self.path.endswith('/generate'):
            # Health check endpoint
            response = json.dumps({'status': 'healthy'}).encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_cors_headers()
            self.send_header('Content-Length', str(len(response)))
            self.end_headers()
            self.wfile.write(response)
        else:
            self.send_error_response(404, 'Not Found')
    
    def send_error_response(self, status_code, message, details=None):
        """Send error response with JSON body"""
        error_data = {'error': message}
        if details:
            error_data['details'] = str(details)
        # Log error for debugging (Vercel will capture this)
        print(f"ERROR [{status_code}]: {message}")
        if details:
            print(f"Details: {details}")
        error_response = json.dumps(error_data).encode('utf-8')
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        self.send_cors_headers()
        self.send_header('Content-Length', str(len(error_response)))
        self.end_headers()
        self.wfile.write(error_response)
    
    def log_message(self, format, *args):
        """Suppress default logging"""
        pass
