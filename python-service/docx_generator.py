from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import io
import json
import os
from datetime import datetime
import re

app = Flask(__name__)
CORS(app)

# Russian number spelling dictionaries
ONES = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
ONES_FEM = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
TENS = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
TEENS = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать', 'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
HUNDREDS = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
THOUSANDS = ['', 'тысяча', 'тысячи', 'тысяч']
MILLIONS = ['', 'миллион', 'миллиона', 'миллионов']
RUBLES = ['рубль', 'рубля', 'рублей']
KOPECKS = ['копейка', 'копейки', 'копеек']
WORKDAYS = ['рабочий день', 'рабочих дня', 'рабочих дней']

MONTHS_RU = ['', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']

def format_number_russian(num):
    """Format number with Russian formatting: space as thousands separator, comma as decimal"""
    num_str = f"{num:.2f}"
    parts = num_str.split('.')
    integer_part = parts[0]
    decimal_part = parts[1] if len(parts) > 1 else '00'
    
    # Add space as thousands separator
    integer_formatted = ''
    for i, digit in enumerate(reversed(integer_part)):
        if i > 0 and i % 3 == 0:
            integer_formatted = ' ' + integer_formatted
        integer_formatted = digit + integer_formatted
    
    return f"{integer_formatted},{decimal_part}"

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def set_font_times_new_roman(run, size=12, bold=False, italic=False):
    """Set font to Times New Roman with specified size and style"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def extract_order_number(order_id):
    """Extract number from order ID (e.g., 'order-27193' -> '27193', 'KP-2025-193' -> '193')"""
    if not order_id:
        return ''
    order_str = str(order_id)
    
    # If order ID contains hyphens, get the last part after the last hyphen
    if '-' in order_str:
        parts = order_str.split('-')
        last_part = parts[-1]
        # Extract digits from the last part only
        digits_only = re.sub(r'\D', '', last_part)
        if digits_only:
            return digits_only
    
    # If no hyphens, extract all digits
    digits_only = re.sub(r'\D', '', order_str)
    if digits_only:
        return digits_only
    
    return order_str

def get_declension(num, forms):
    """Get proper Russian declension based on number"""
    num = abs(num) % 100
    num1 = num % 10
    if num > 10 and num < 20:
        return forms[2]
    if num1 > 1 and num1 < 5:
        return forms[1]
    if num1 == 1:
        return forms[0]
    return forms[2]

def spell_number_russian(num, feminine=False):
    """Convert number to Russian words"""
    if num == 0:
        return 'ноль'
    
    result = []
    num = int(num)
    
    # Millions
    if num >= 1000000:
        millions = num // 1000000
        result.append(spell_number_russian(millions, False))
        result.append(get_declension(millions, MILLIONS))
        num = num % 1000000
    
    # Thousands
    if num >= 1000:
        thousands = num // 1000
        if thousands == 1:
            result.append('одна')
        elif thousands == 2:
            result.append('две')
        else:
            result.append(spell_number_russian(thousands, True))
        result.append(get_declension(thousands, THOUSANDS))
        num = num % 1000
    
    # Hundreds
    if num >= 100:
        hundreds = num // 100
        result.append(HUNDREDS[hundreds])
        num = num % 100
    
    # Tens and ones
    if num >= 20:
        tens = num // 10
        result.append(TENS[tens])
        num = num % 10
    
    if num >= 10:
        result.append(TEENS[num - 10])
        num = 0
    
    if num > 0:
        if feminine:
            result.append(ONES_FEM[num])
        else:
            result.append(ONES[num])
    
    return ' '.join(result)

def spell_money_russian(amount):
    """Convert money amount to Russian words"""
    rubles = int(amount)
    kopecks = int(round((amount - rubles) * 100))
    
    rubles_text = spell_number_russian(rubles, False)
    rubles_form = get_declension(rubles, RUBLES)
    
    result = f"{rubles_text.capitalize()} {rubles_form}"
    
    if kopecks > 0:
        kopecks_text = spell_number_russian(kopecks, True)
        kopecks_form = get_declension(kopecks, KOPECKS)
        result += f" {kopecks_text} {kopecks_form}"
    else:
        result += " 00 копеек"
    
    return result

def spell_workdays_russian(days):
    """Convert workdays number to Russian words"""
    days_text = spell_number_russian(days, False)
    days_form = get_declension(days, WORKDAYS)
    return f"{days_text.capitalize()} {days_form}"

def format_date_full_russian(date_str):
    """Format date to full Russian format"""
    try:
        # Try parsing different date formats
        if ' ' in date_str:
            # Format like "Jan 15, 2025"
            dt = datetime.strptime(date_str, '%b %d, %Y')
        elif '-' in date_str:
            # Format like "2025-01-15"
            dt = datetime.strptime(date_str.split()[0], '%Y-%m-%d')
        else:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
        
        day = dt.day
        month = MONTHS_RU[dt.month]
        year = dt.year
        
        return f"{day} {month} {year} г."
    except:
        return date_str

def set_cell_border(cell, **kwargs):
    """Set border for a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data['val'])
            element.set(qn('w:sz'), str(edge_data.get('sz', 6)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            borders.append(element)
    tcPr.append(borders)

def add_company_header(doc, company_data, locale='ru-RU', available_width=Inches(7.5)):
    """Add company information header in Russian format using a table with border"""
    # Create header table
    header_table = doc.add_table(rows=1, cols=1)
    header_table.columns[0].width = available_width
    
    header_cell = header_table.cell(0, 0)
    
    # Line 1: ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ
    line1_para = header_cell.paragraphs[0]
    line1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1_run = line1_para.add_run("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ")
    set_font_times_new_roman(line1_run, size=12, bold=True, italic=True)
    
    # Line 2: «МК СЕРВИС»
    line2_para = header_cell.add_paragraph()
    line2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_name = company_data.get('name', '')
    line2_run = line2_para.add_run(f"«{company_name}»")
    set_font_times_new_roman(line2_run, size=24, bold=True, italic=True)
    
    # Line 3: Address
    if company_data.get('address'):
        line3_para = header_cell.add_paragraph()
        line3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line3_run = line3_para.add_run(company_data['address'])
        set_font_times_new_roman(line3_run, size=9, bold=False, italic=False)
    
    # Line 4: Bank details
    if locale and locale.startswith('ru'):
        bank_parts = []
        if company_data.get('bankName'):
            bank_parts.append(company_data['bankName'])
        if company_data.get('bankAccount'):
            bank_parts.append(f"р/с {company_data['bankAccount']}")
        if company_data.get('correspondentAccount'):
            bank_parts.append(f"к/с {company_data['correspondentAccount']}")
        if company_data.get('bankBik'):
            bank_parts.append(f"БИК {company_data['bankBik']}")
        
        if bank_parts:
            line4_para = header_cell.add_paragraph()
            line4_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line4_run = line4_para.add_run(', '.join(bank_parts))
            set_font_times_new_roman(line4_run, size=9, bold=False, italic=False)
    
    # Line 5: ИНН and КПП
    if locale and locale.startswith('ru'):
        company_parts = []
        if company_data.get('inn'):
            company_parts.append(f"ИНН {company_data['inn']}")
        if company_data.get('kpp'):
            company_parts.append(f"КПП {company_data['kpp']}")
        
        if company_parts:
            line5_para = header_cell.add_paragraph()
            line5_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line5_run = line5_para.add_run(', '.join(company_parts))
            set_font_times_new_roman(line5_run, size=9, bold=False, italic=False)
    
    # Set bottom border for the table cell
    set_cell_border(header_cell, bottom={'val': 'single', 'sz': 6, 'color': '000000'})
    
    doc.add_paragraph()

def add_document_header(doc, order_data, client_data, doc_type, locale='ru-RU', available_width=Inches(7.5)):
    """Add document header (Смета № or КП №)"""
    if doc_type == 'invoice':
        doc_label = "Смета №"
    else:  # po
        doc_label = "КП №"
    
    # Always use order ID to extract number
    order_id = order_data.get('id', '')
    doc_number = extract_order_number(order_id)
    
    # Use current date
    current_date = datetime.now()
    if locale and locale.startswith('ru'):
        date_formatted = format_date_full_russian(current_date.strftime('%Y-%m-%d'))
    else:
        date_formatted = current_date.strftime('%Y-%m-%d')
    
    # Create table for alignment - use full width
    header_table = doc.add_table(rows=2, cols=2)
    # Extract inch value from Inches object
    width_val = available_width.inches if hasattr(available_width, 'inches') else float(available_width)
    header_table.columns[0].width = Inches(width_val * 0.65)  # 65% for left content
    header_table.columns[1].width = Inches(width_val * 0.35)  # 35% for date
    
    # Left column: Document number and client
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(f"{doc_label} {doc_number}")
    set_font_times_new_roman(left_run, size=12, bold=False, italic=False)
    
    # Client name on next line
    client_name = client_data.get('company') or client_data.get('name') or 'Клиента'
    left_para2 = left_cell.add_paragraph()
    left_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run2 = left_para2.add_run(f"для {client_name}")
    set_font_times_new_roman(left_run2, size=12, bold=False, italic=False)
    
    # Right column: Date (right-aligned, spanning both rows)
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(date_formatted)
    set_font_times_new_roman(right_run, size=12, bold=False, italic=False)
    
    # Merge right cell to span both rows
    right_cell2 = header_table.cell(1, 1)
    right_cell.merge(right_cell2)
    
    doc.add_paragraph()

    # Order title (centered, bold)
    if order_data.get('orderTitle'):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(order_data['orderTitle'])
        set_font_times_new_roman(title_run, size=12, bold=True, italic=False)
    
    doc.add_paragraph()

def add_work_description(doc, jobs_data, order_data, locale='ru-RU', work_days=30, available_width=Inches(7.5)):
    """Add work description with table view for jobs"""
    if not jobs_data:
        doc.add_paragraph("Нет работ")
        return
    
    # Create jobs table with 5 columns - use full available width
    jobs_table = doc.add_table(rows=1, cols=5)
    jobs_table.style = 'Table Grid'
    # Extract inch value from Inches object
    width_val = available_width.inches if hasattr(available_width, 'inches') else float(available_width)
    # Proportional widths: № (5%), Наименование (50%), Кол-во (10%), Цена за единицу (15%), Стоимость (20%)
    jobs_table.columns[0].width = Inches(width_val * 0.05)  # №
    jobs_table.columns[1].width = Inches(width_val * 0.50)  # Наименование
    jobs_table.columns[2].width = Inches(width_val * 0.10)  # Кол-во
    jobs_table.columns[3].width = Inches(width_val * 0.15)  # Цена за единицу
    jobs_table.columns[4].width = Inches(width_val * 0.20)  # Стоимость
    
    # Header row
    header_cells = jobs_table.rows[0].cells
    headers = ['№', 'Наименование', 'Кол-во', 'Цена за единицу', 'Стоимость']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Set grey background
        set_cell_shading(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font_times_new_roman(run, size=12, bold=True, italic=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add job rows
    for idx, job in enumerate(jobs_data, start=1):
        row_cells = jobs_table.add_row().cells
        
        # №
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run0 = row_cells[0].paragraphs[0].add_run(str(idx))
        set_font_times_new_roman(run0, size=12, bold=False, italic=False)
        
        # Наименование
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = row_cells[1].paragraphs[0].add_run(job.get('name', ''))
        set_font_times_new_roman(run1, size=12, bold=False, italic=False)
        
        # Кол-во (without unit)
        job_qty = job.get('qty', '0')
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = row_cells[2].paragraphs[0].add_run(job_qty)
        set_font_times_new_roman(run2, size=12, bold=False, italic=False)
        
        # Цена за единицу (after markup) - calculate from lineTotal / quantity
        job_total = float(job.get('lineTotal', '0.00'))
        job_qty_float = float(job.get('qty', '1'))
        job_price_after_markup = job_total / job_qty_float if job_qty_float > 0 else 0.00
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run3 = row_cells[3].paragraphs[0].add_run(format_number_russian(job_price_after_markup))
        set_font_times_new_roman(run3, size=12, bold=False, italic=False)
        
        # Стоимость
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run4 = row_cells[4].paragraphs[0].add_run(format_number_russian(job_total))
        set_font_times_new_roman(run4, size=12, bold=False, italic=False)
    
    doc.add_paragraph()

    # Cost line with Russian number formatting and spelling on same line
    total = float(order_data.get('total', '0.00'))
    cost_para = doc.add_paragraph()
    if locale and locale.startswith('ru'):
        total_spelled = spell_money_russian(total)
        cost_run = cost_para.add_run(f"Стоимость работ по заказу составляет: {format_number_russian(total)} руб. ({total_spelled}) Без НДС.")
    else:
        cost_run = cost_para.add_run(f"Стоимость работ по заказу составляет: {format_number_russian(total)} руб.")
    set_font_times_new_roman(cost_run, size=12, bold=False, italic=False)
    
    # Add new line after cost line
    doc.add_paragraph()
    
    # Work completion deadline
    if locale and locale.startswith('ru'):
        days_spelled = spell_workdays_russian(work_days)
        deadline_para = doc.add_paragraph()
        deadline_run = deadline_para.add_run(f"Срок выполнения работ – {work_days} ({days_spelled}) рабочих дней с момента внесения предоплаты и подписания сметы.")
        set_font_times_new_roman(deadline_run, size=12, bold=False, italic=False)
        
        # Note about deadline extension
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("Сроки могут быть увеличены по согласованию сторон, в случае проведения дополнительных работ.")
        set_font_times_new_roman(note_run, size=12, bold=False, italic=False)
        
        # Add 2 new lines after note
        doc.add_paragraph()
        doc.add_paragraph()
    
    doc.add_paragraph()

def add_footer_section(doc, company_data, doc_type, locale='ru-RU', available_width=Inches(7.5)):
    """Add footer section with signature and contact info"""
    if not (locale and locale.startswith('ru')):
        return
    
    doc.add_paragraph()
    
    if doc_type == 'invoice':
        # For invoices: 3-column layout with signatures - use full width
        footer_table = doc.add_table(rows=4, cols=3)
        # Extract inch value from Inches object
        width_val = available_width.inches if hasattr(available_width, 'inches') else float(available_width)
        footer_table.columns[0].width = Inches(width_val / 3)  # Исполнитель
        footer_table.columns[1].width = Inches(width_val / 3)  # Заказчик
        footer_table.columns[2].width = Inches(width_val / 3)  # Contact info
        
        # Row 1: Headers
        left_header = footer_table.cell(0, 0)
        left_header.text = "Исполнитель"
        for paragraph in left_header.paragraphs:
            for run in paragraph.runs:
                set_font_times_new_roman(run, size=12, bold=True, italic=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        right_header = footer_table.cell(0, 1)
        right_header.text = "Заказчик"
        for paragraph in right_header.paragraphs:
            for run in paragraph.runs:
                set_font_times_new_roman(run, size=12, bold=True, italic=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Contact header in right column
        contact_header = footer_table.cell(0, 2)
        contact_header_para = contact_header.paragraphs[0]
        contact_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_header_run = contact_header_para.add_run("С уважением,")
        set_font_times_new_roman(contact_header_run, size=12, bold=False, italic=False)
        
        # Director info immediately after (no empty line) - add to same paragraph
        if company_data.get('directorName'):
            director_run1 = contact_header_para.add_run("\nТехнический директор")
            set_font_times_new_roman(director_run1, size=12, bold=False, italic=False)
            director_run2 = contact_header_para.add_run(f"\n{company_data.get('legalName', company_data.get('name', ''))}")
            set_font_times_new_roman(director_run2, size=12, bold=False, italic=False)
            director_run3 = contact_header_para.add_run(f"\n{company_data['directorName']}")
            set_font_times_new_roman(director_run3, size=12, bold=False, italic=False)
        
        # Row 2: Signature lines
        left_sig = footer_table.cell(1, 0)
        left_sig.text = "_________________"
        for para in left_sig.paragraphs:
            for run in para.runs:
                set_font_times_new_roman(run, size=12, bold=False, italic=False)
        
        right_sig = footer_table.cell(1, 1)
        right_sig.text = "_________________"
        for para in right_sig.paragraphs:
            for run in para.runs:
                set_font_times_new_roman(run, size=12, bold=False, italic=False)
        
        # Contact info in right column (bottom right)
        contact_cell = footer_table.cell(1, 2)
        contact_para = contact_cell.paragraphs[0]
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if company_data.get('email'):
            contact_run1 = contact_para.add_run(f"Email: {company_data['email']}")
            set_font_times_new_roman(contact_run1, size=12, bold=False, italic=False)
        if company_data.get('phone'):
            if company_data.get('email'):
                contact_run2 = contact_para.add_run(f"\nТел: {company_data['phone']}")
            else:
                contact_run2 = contact_para.add_run(f"Тел: {company_data['phone']}")
            set_font_times_new_roman(contact_run2, size=12, bold=False, italic=False)
        
        # Row 3: Name lines
        left_name = footer_table.cell(2, 0)
        left_name.text = "_________________"
        for para in left_name.paragraphs:
            for run in para.runs:
                set_font_times_new_roman(run, size=12, bold=False, italic=False)
        
        right_name = footer_table.cell(2, 1)
        right_name.text = "_________________"
        for para in right_name.paragraphs:
            for run in para.runs:
                set_font_times_new_roman(run, size=12, bold=False, italic=False)
        
        # Row 4: Empty
        footer_table.cell(3, 0).text = ""
        footer_table.cell(3, 1).text = ""
        footer_table.cell(3, 2).text = ""
    else:
        # For PO: 2-column layout - left: director info, right: contact info - use full width
        footer_table = doc.add_table(rows=1, cols=2)
        # Extract inch value from Inches object
        width_val = available_width.inches if hasattr(available_width, 'inches') else float(available_width)
        footer_table.columns[0].width = Inches(width_val * 0.55)  # Director info (55%)
        footer_table.columns[1].width = Inches(width_val * 0.45)  # Contact info (45%)
        
        # Left column: Director info
        left_cell = footer_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run1 = left_para.add_run("С уважением,")
        set_font_times_new_roman(left_run1, size=12, bold=False, italic=False)
        if company_data.get('directorName'):
            left_run2 = left_para.add_run("\n\nТехнический директор")
            set_font_times_new_roman(left_run2, size=12, bold=False, italic=False)
            left_run3 = left_para.add_run(f"\n{company_data.get('legalName', company_data.get('name', ''))}")
            set_font_times_new_roman(left_run3, size=12, bold=False, italic=False)
            left_run4 = left_para.add_run(f"\n{company_data['directorName']}")
            set_font_times_new_roman(left_run4, size=12, bold=False, italic=False)
        
        # Right column: Contact info (bottom right)
        right_cell = footer_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if company_data.get('email'):
            right_run1 = right_para.add_run(f"Email: {company_data['email']}")
            set_font_times_new_roman(right_run1, size=12, bold=False, italic=False)
        if company_data.get('phone'):
            if company_data.get('email'):
                right_run2 = right_para.add_run(f"\nТел: {company_data['phone']}")
            else:
                right_run2 = right_para.add_run(f"Тел: {company_data['phone']}")
            set_font_times_new_roman(right_run2, size=12, bold=False, italic=False)

def generate_document(data, doc_type):
    """Generate DOCX document from data"""
    doc = Document()
    
    # Set narrow margins (0.5 inches on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Available width with narrow margins: 8.5" - 0.5" - 0.5" = 7.5"
    AVAILABLE_WIDTH = Inches(7.5)
    
    # Set default font to Times New Roman, size 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set line spacing to 0 (single line spacing)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0  # Single line spacing (0 would be no spacing)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    locale = data.get('locale', 'ru-RU')
    company_data = data.get('company', {})
    client_data = data.get('client', {})
    order_data = data.get('order', {})
    jobs_data = data.get('jobs', [])
    work_days = data.get('workCompletionDays', 30)
    
    # 1. Company header with banking
    add_company_header(doc, company_data, locale, AVAILABLE_WIDTH)
    
    # 2. Document header (Смета № or КП №)
    add_document_header(doc, order_data, client_data, doc_type, locale, AVAILABLE_WIDTH)
    
    # 3. Work description with table view
    add_work_description(doc, jobs_data, order_data, locale, work_days, AVAILABLE_WIDTH)
    
    # 4. Footer with signature and contact (3-column layout)
    add_footer_section(doc, company_data, doc_type, locale, AVAILABLE_WIDTH)
    
    return doc

@app.route('/generate', methods=['POST'])
def generate_docx():
    """API endpoint to generate DOCX documents"""
    try:
        data = request.get_json()
        doc_type = data.get('type', 'po')
        
        # Generate document
        doc = generate_document(data, doc_type)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Return file
        filename = f"{doc_type}-{data.get('order', {}).get('id', 'document')}.docx"
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    debug = os.environ.get('FLASK_DEBUG', 'True').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug)
